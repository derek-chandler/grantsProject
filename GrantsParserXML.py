import xml.etree.ElementTree as et
import html as html
import GrantDownloader
from tkinter import *
from tkinter import messagebox
from tkcalendar import DateEntry
import os
import datetime
import word
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# dictionary of agencies using agency code as key
# these were all the agencies in the search function for Grants.gov
# I added 'N/A' to the list to make sure that if we did not have a match, we would still have a key for it
agencyDictionary = {'USAID': 'Agency for International Development',
                    'AC': 'AmeriCorps',
                    'USDA': 'Department of Agriculture',
                    'DOC': 'Department of Commerce',
                    'DOE': 'Department of Energy',
                    'DOD': 'Department of Defense',
                    'ED': 'Department of Education',
                    'PAMS': 'Department of Health and Human Services',
                    'HHS': 'Department of Health and Human Services',
                    'DHS': 'Department of Homeland Security',
                    'HUD': 'Department of Housing and Urban Development',
                    'USDOJ': 'Department of Justice',
                    'DOL': 'Department of Labor',
                    'DOS': 'Department of State',
                    'DOI': 'Department of the Interior',
                    'USDOT': 'Department of the Treasury',
                    'DOT': 'Department of Transportation',
                    'VA': 'Department of Veterans Affairs',
                    'EPA': 'Environmental Protection Agency',
                    'GCERC': 'Gulf Coast Ecosystem Restoration Council',
                    'IMLS': 'Institute of Museum and Library Services',
                    'MCC': 'Millennium Challenge Corportation',
                    'NASA': 'National Aeronautics and Space Administration',
                    'NARA': 'National Archives and Records Administration',
                    'NEA': 'National Endowment for the Arts',
                    'NEH': 'National Endowment for the Humanities',
                    'NSF': 'National Science Foundation',
                    'NRC': 'National Resource Conservation Council',
                    'SBA': 'Small Business Administration',
                    'SSA': 'Social Security Administration',
                    'N/A': 'Other Agencies'}

# String that saves the http portion of the XML tags so that we don't have to keep typing it out,
# e.g. referencing the tag '{http://apply.grants.gov/system/OpportunityDetail-V1.0}OpportunityTitle'
# becomes just linkString+'OpportunityTitle'
linkString = '{http://apply.grants.gov/system/OpportunityDetail-V1.0}'

# ********************************************DEF*****************************************************************

# Convert our dates into a better looking format with slashes, MM/DD/YYYY


def dateConversion(date):
    newDate = date[:2] + "/" + date[2:4] + "/" + date[4:]
    return newDate

# Convert our dates to format similar to January 01, 2021


def dateStringVersion(date):
    newDate = ''
    if (date != 'N/A'):

        str(date)
        monthList = ['January ', 'February ', 'March ', 'April ', 'May ', 'June ',
                     'July ', 'August ', 'September ', 'October ', 'November ', 'December ']
        monthNum = int(date[:2])
        temp = monthList[monthNum - 1]
        newDate = temp + date[2:4] + ", " + date[4:]
    else:
        newDate = 'N/A'
    return newDate


# input a string to add commas
# example input : 10000000
# example output: 10,000,000
def addCommasAndDollarSign(amountStr):
    # check if string is a number, if not return the string back
    if not amountStr.isnumeric():
        return amountStr
    else:
        return "${:,}".format(int(amountStr))


# convert our dates into a year, month, day hierarchy so that earlier dates are natrually smaller numbers (strings in this case) than later dates
def dateHierarchyForm(date):
    newDate = date[4:] + date[:4]
    return newDate


# this function takes an agency code and returns the name of the agency
def generateAgencyName(agencyCode):
    # agency codes have dashes to separate the information.
    if '-' in agencyCode:
        # we only want the first part of the agency code which identifies the agency itself
        agencyCode = agencyCode.split('-')[0]

    # if the agency is in the list of agencies, we can use the agency code to get the name
    if agencyCode in agencyDictionary:
        agencyCode = agencyDictionary[agencyCode]   # of the agency

    else:
        # if the agency is not in the list of agencies, we will use the string 'Other Agencies'
        agencyCode = 'Other Agencies'

    return agencyCode


# create table of unique agencies (want to improve this, lot of redudancies related
# to multiple instances of the same agency with additional information)
def tableOfContents(listA, agency):

    if agency in listA:                     # if the agency is already in the list, we don't need to add it again
        return listA
    else:
        # if the agency is not in the list, we add it to the list
        listA.append(agency)

    return listA


# function to return an attribute of interest from a given opportunity
def getOpportunityInfo(opportunity, attribute):
    #
    # Store each text of an attribute as a string, or store as 'N/A' if none exist
    myInfo = getattr(opportunity.find(linkString + attribute), 'text', 'N/A')
    # at given attribute location
    return myInfo


# ********************************************MAIN Object*****************************************************************
class Grant:

    def __init__(self, agencyCode, agencyName, opportunityTitle, postDate, dueDate, numAwards,
                 totalFunding, awardCeiling, awardFloor, oppNumber, description, eligApplicants='N/A'):

        self.agencyCode = agencyCode
        self.distinctAgency = generateAgencyName(agencyCode)
        self.agencyName = agencyName
        self.opportunityTitle = opportunityTitle
        self.postDate = postDate
        self.dueDate = dueDate
        self.numAwards = numAwards
        self.totalFunding = addCommasAndDollarSign(totalFunding)
        self.awardCeiling = addCommasAndDollarSign(awardCeiling)
        self.awardFloor = addCommasAndDollarSign(awardFloor)
        self.oppNumber = oppNumber
        self.description = description
        self.eligApplicants = eligApplicants

# ********************************************MAIN FUNCTIONS*****************************************************************

def printGrant(grant):
    print("Agency name:                     " + grant.agencyName)
    print("Opportunity title:               " + grant.opportunityTitle)
    print("Post date:                       " +
          dateStringVersion(grant.postDate))
    print("Due date:                        " +
          dateStringVersion(grant.dueDate))
    print("Expected Number of awards:       " + grant.numAwards)
    print("Estimated total program funding: " + addCommasAndDollarSign(grant.totalFunding))
    print("Award Ceiling:                   " + addCommasAndDollarSign(grant.awardCeiling))
    print("Award floor:                     " + addCommasAndDollarSign(grant.awardFloor))
    print("Funding opportunity number:      " + grant.oppNumber)
    print()
    print("Purpose: " + grant.description)
    print()
    print("Eligible applicants: " + grant.eligApplicants)

    print()


# function to create a dictionary using the distinctAgency as a key and the grants as values
def grantDictionaryAdd(grantDictionary, grant):
    grantDictionary.setdefault(grant.distinctAgency, []).append(grant)
    return grantDictionary

# ********************************DRIVER_CODE****************************************************************************


'''
Trying to look into more efficient parsing, this site seems to have a better, space efficient parsing method I need to explore:
https://newbedev.com/efficient-way-to-iterate-through-xml-elements
'''

# Parse our file and store it in a tree
mytree = et.parse(GrantDownloader.get())

# --------------------------- UI BEGIN ---------------------------

# Set today's date and set last week's date
today = datetime.date.today()
last_week = today - datetime.timedelta(days=7)

# Basic Settings, window title / size
root = Tk()
root.title('US Government Grant Report Tool')
if os.name == 'posix':
    root.iconbitmap('@resource/tux.xbm')
else:
    root.iconbitmap('resource/icon.ico')
root.geometry("500x320")

# Sets layout of modules
top = Frame(root)
bottom = Frame(root, width=100)
top.pack(side=TOP)
bottom.pack(side=BOTTOM, fill=None, expand=False)

# Sets padding and text of UI Label
my_toplabel = Label(root, text="Please select a date range")
my_toplabel.pack(pady=10, in_=top)

# Set and post date entry fields (first set to 1 week past, second to current date)
calone = DateEntry(root, width=12, background='darkblue',
                   foreground='white', borderwidth=2, year=last_week.year, month=last_week.month, day=last_week.day)
calone.pack(in_=top, side=LEFT, padx=20, pady=10)

caltwo = DateEntry(root, width=12, background='darkblue',
                   foreground='white', borderwidth=2)
caltwo.pack(in_=top, side=RIGHT, padx=20, pady=10)


# Function for collecting date from calendar
def grab_date():
    global userdateone, userdatetwo
    userdateone = calone.get_date()
    userdatetwo = caltwo.get_date()
    if userdateone > userdatetwo:
        messagebox.showerror("Improper Date Range", "Please ensure your first date is before your second date")
    else: root.destroy()


# Button Grabs selected dates then closes if userdateone > userdatetwo
my_button = Button(root, text="Confirm", activebackground='gray', command=grab_date)
my_button.pack(pady=10, padx=10, in_=bottom, side=RIGHT)

# Location for date to post
my_label = Label(root, text="")
my_label.pack(pady=10)

# loopy boi
root.mainloop()

# Convert datetime object to string for comparison
dateRangeOne = userdateone.strftime("%Y%m%d")
dateRangeTwo = userdateone.strftime("%Y%m%d")
print(dateRangeTwo, dateRangeOne)

# --------------------------- UI END ---------------------------

# Store the root element of this file
myroot = mytree.getroot()

# Count the number of grants that we have chosen to print,
count = 0

# Declare our list to store agency names
agencyList = []
grantDictionary = {}

# Check each grant opportunity in our tree
for opportunity in myroot:

    # Get the postdate first
    # getattr(opportunity.find(linkString + 'PostDate'), 'text', 'N/A')
    postDate = getOpportunityInfo(opportunity, 'PostDate')
    # Set the desired earliest date
    # Finds grants of date date range selected in UI
    if dateRangeOne <= dateHierarchyForm(postDate) <= dateRangeTwo:

        # print('************************************************************************************************************************')
        # print()

        # Store each text of qualifying grants as a string, or store as 'N/A' if none exist
        grant = Grant(
            agencyCode=getOpportunityInfo(opportunity, 'AgencyCode'),
            agencyName=getOpportunityInfo(opportunity, 'AgencyName'),
            opportunityTitle=html.unescape(
                getOpportunityInfo(opportunity, 'OpportunityTitle')),
            postDate=getOpportunityInfo(opportunity, 'PostDate'),
            dueDate=getOpportunityInfo(opportunity, 'DueDate'),
            numAwards=getOpportunityInfo(opportunity, 'NumberOfAwards'),
            totalFunding=getOpportunityInfo(opportunity, 'TotalFunding'),
            awardCeiling=getOpportunityInfo(opportunity, 'AwardCeiling'),
            awardFloor=getOpportunityInfo(opportunity, 'AwardFloor'),
            oppNumber=getOpportunityInfo(opportunity, 'OpportunityNumber'),
            description=html.unescape(
                getOpportunityInfo(opportunity, 'Description')),
            eligApplicants=getOpportunityInfo(
                opportunity, 'EligibleApplicants')
        )
        tableOfContents(agencyList, grant.distinctAgency)
        # Create a grant object
        grantDictionary = grantDictionaryAdd(grantDictionary, grant)
        # printGrant(grant)

        # Count the selected grant
        count += 1

# print('**********************************************************************************************************')
# # Print out the number of grants that qualified. I used this to check to make sure pruning was happening
# print("Number of grants", count)

# # sort the list of agencies
# print()
agencyList.sort()

# print the list of agencies
# print('Table of Contents')
# print('----------------------------------------------------------------------')
# for x in agencyList:
#     print(x)

# # Using the grantDictionary, print out each grant for Department of Education key
# print('----------------------------------------------------------------------')
# print()
# print('ALL DEPARTMENT OF EDUCATION GRANTS')
# print()
# for gr in grantDictionary['Department of Education']:
#     print('**********************************************************************************************************')
#     printGrant(gr)

#! Opens the word templet file
doc = docx.Document("templet.docx")

bookmark_list = list()

#! change text in paragraph 9 to the number of grants
doc.paragraphs[9].text = str(date.today().strftime("%B %d, %Y"))

#! Set pointer to equal paragraph 11
paracount = 11
pointer = doc.paragraphs[paracount]

#! Build the table of contents first
for agency in agencyList:
    word.insert_paragraph_after(pointer, agency)
    paracount += 1
    pointer = doc.paragraphs[paracount]
    
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line.add_run("\nGrants\n")
run.bold = True
font = run.font
font.size = Pt(22)
font.name = 'Times New Roman'
font.underline = True


#! This prints generates the bookmarks
for agency  in agencyList:

    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.0
    word.add_bookmark(paragraph, agency, f"{agency} bookmark")
    bookmark_list.append(paragraph)

    grantDictionary[agency].sort(key=lambda x: x.dueDate)

    # Loop over each grant in the dictionary
    grant_list = grantDictionary.get(agency)
    for i in grant_list:
        paragraph = doc.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.0

        paragraph.add_run(f"\nAgency Name: {i.opportunityTitle}").bold = True
        paragraph.add_run(f"\nOpportunity Title: {i.opportunityTitle}").bold = True
        paragraph.add_run(f"\nPost Date: {i.postDate}").bold = True
        paragraph.add_run(f"\nDue Date {i.dueDate}").bold = True
        paragraph.add_run(f"\nExpected Number of awards: {i.numAwards}").bold = True
        paragraph.add_run(f"\nEstimated total program funding: {i.totalFunding}").bold = True
        paragraph.add_run(f"\nAward Ceiling: {i.awardCeiling}").bold = True
        paragraph.add_run(f"\nAward Floor: {i.awardFloor}").bold = True
        paragraph.add_run(f"\nFunding Opportunity Number: {i.oppNumber}").bold = True

        run = paragraph.add_run(f"\n\nPurpose: {i.description}")
        font = run.font
        font.size = Pt(12)
        font.italic = True
        font.name = 'Times New Roman'


        paragraph.add_run(f"\n\nEligible Applicants: {i.eligApplicants}").bold = True
        paragraph.add_run(f"\n\n")


for index, bookmark in enumerate(bookmark_list):
    """_summary_

    Args:
        bookmarkinenumerate (_type_): _description_
    """
    word.add_link(bookmark, agencyList[index], f"{agencyList[index]} link")

    # # table_of_content = doc.paragraphs[11].add_run(f"{agency}")
    # table_of_content = doc.paragraphs[11]
    # toc = doc.add_paragraph(agency)
    # table_of_content.insert_paragraph_after(toc)



# for i in grantDictionary["Agency for International Development"]:

#     print(i.awardCeiling)

doc.save("sample.docx")
